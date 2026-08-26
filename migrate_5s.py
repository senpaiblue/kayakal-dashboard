import os
import shutil
import re
from pathlib import Path
import pandas as pd
from docx import Document

# Paths
excel_in = 'Data/5S Model area.xlsx'
dt_path = 'assets/5s/DT.xlsx'
base_assets = Path('assets/5s')

# --- STEP 1: PARSE AND REBUILD DT.xlsx ---
print("--- Step 1: Parsing Data/5S Model area.xlsx ---")
if not os.path.exists(excel_in):
    raise FileNotFoundError(f"Missing {excel_in}")

# Read the excel sheet
df_new = pd.read_excel(excel_in, sheet_name='5S Implementation - Nominat (2)', header=None)
df_new.columns = df_new.iloc[1]
df_new = df_new.iloc[2:].reset_index(drop=True)
df_new.columns = [str(col).strip() for col in df_new.columns]

# Forward fill department and zone
df_new[['Department', 'Zone']] = df_new[['Department', 'Zone']].ffill()

# Clean values
df_new['Department'] = df_new['Department'].astype(str).str.strip()
df_new['5S Model area'] = df_new['5S Model area'].astype(str).str.strip()
df_new['Team Leader'] = df_new['Team Leader'].astype(str).str.strip()

# Group by (Department, 5S Model area) and aggregate Team Leaders with ' / '
df_grouped = df_new.groupby(['Department', '5S Model area'], as_index=False).agg({
    'Team Leader': lambda series: " / ".join(sorted(set(series.dropna())))
})

# Read existing passwords from current DT.xlsx to preserve them
existing_pwds = {}
if os.path.exists(dt_path):
    df_old_dt = pd.read_excel(dt_path)
    for idx, r in df_old_dt.iterrows():
        dept_val = str(r['department']).strip()
        pwd_val = str(r['password']).strip()
        if dept_val and pwd_val and pwd_val.lower() != 'nan':
            existing_pwds[dept_val.lower()] = pwd_val

# Rebuild rows for new DT.xlsx
merged_rows = []
for idx, r in df_grouped.iterrows():
    dept = r['Department']
    # Sanitize model name - replace slashes with hyphens to make safe on Windows and prevent nested folder bugs
    model = re.sub(r'\s*/\s*', ' - ', str(r['5S Model area'])).strip()
    leader = r['Team Leader']
    
    # Password matching (case insensitive department lookup)
    pwd = existing_pwds.get(dept.lower(), 'JSW@2026')
    
    merged_rows.append({
        'department': dept,
        'model': model,
        'leader': leader,
        'password': pwd
    })

df_dt_new = pd.DataFrame(merged_rows)
# Save backup of DT.xlsx first
if os.path.exists(dt_path):
    shutil.copy2(dt_path, dt_path + '.bak')
df_dt_new.to_excel(dt_path, index=False)
print(f"Rebuilt DT.xlsx successfully with {len(df_dt_new)} unique model areas.")

# --- STEP 2: SAFE FILE MOVES & MERGES ---
print("\n--- Step 2: Running folder migration & merges ---")

def merge_docx_files(src_path, dest_path):
    """Safely merges paragraphs of two docx files into one."""
    print(f"Merging Word Document: {src_path} into {dest_path}")
    doc_src = Document(src_path)
    doc_dest = Document(dest_path)
    
    # Add separator and copy paragraphs
    doc_dest.add_paragraph("\n--- Merged Audit Summary ---")
    for p in doc_src.paragraphs:
        doc_dest.add_paragraph(p.text)
        
    doc_dest.save(dest_path)

def safe_move_file(src_file, dest_dir):
    """Moves a file to dest_dir. If filename conflicts, adds suffix to be unique."""
    dest_dir.mkdir(parents=True, exist_ok=True)
    dest_path = dest_dir / src_file.name
    
    if not dest_path.exists():
        shutil.move(str(src_file), str(dest_path))
    else:
        if src_file.suffix.lower() == '.docx':
            # Merge docx files instead of overwriting
            merge_docx_files(str(src_file), str(dest_path))
            os.remove(src_file)  # Clean up source
        else:
            # Rename image to avoid conflict
            base = dest_path.stem
            ext = dest_path.suffix
            counter = 1
            while True:
                new_name = f"{base}_{counter}{ext}"
                new_dest_path = dest_dir / new_name
                if not new_dest_path.exists():
                    shutil.move(str(src_file), str(new_dest_path))
                    break
                counter += 1

def move_folder_contents(src_dir, dest_dir):
    """Moves all files from src_dir to dest_dir safely, then removes src_dir."""
    src_dir = Path(src_dir)
    dest_dir = Path(dest_dir)
    if not src_dir.exists():
        return
    for item in src_dir.iterdir():
        if item.is_file():
            safe_move_file(item, dest_dir)
        elif item.is_dir() and item.name == 'self':
            # Also move self folder contents
            move_folder_contents(item, dest_dir / 'self')
    # Clean up empty directory tree
    try:
        shutil.rmtree(str(src_dir))
    except Exception:
        pass

# Defined folder mappings (old_path_relative -> new_path_relative)
# relative to base_assets
folder_mappings = [
    # Blast Furnace-1,2
    ('Blast Furnace-1,2/Stock house substore', 'Blast Furnace-1,2/BF-2 Stock house substore'),
    ('Blast Furnace-1,2/E&I store', 'Blast Furnace-1,2/BF-2 E&I store'),
    ('Blast Furnace-1,2/blah blah', 'Blast Furnace-1,2/BF-2 Stock house substore'),
    
    # SMS-1
    ('SMS-1/Caster repair area (1 and 2)', 'SMS-1/SMS-1 Caster repair area'),
    
    # Oxygen Plant
    ('Oxygen Plant/Control Room', 'Oxygen Plant/Oxygen Control Room'),
    
    # HSM-2
    ('HSM-2/Sub Store', 'HSM-2/HSM-2 Sub Store'),
    ('HSM-2/Roll shop', 'HSM-2/HSM-2 Roll shop'),
    
    # BRM-1
    ('BRM-1/RollShop', 'BRM-1/BRM-1 RollShop'),
    
    # CRM-1
    ('CRM-1/CGl-2 Exit ECR', 'CRM-1/CRM-1 CGL2 Exit ECR'),
    ('CRM-1/RCL Exit Mechanical Area - Tool Room', 'CRM-1/CRM-1 RCL Exit Mechanical Area Tool Room'),
    ('CRM-1/CGL3', 'CRM-1/CRM-1 SPM-ECR'),
    ('CRM-1/SPM-ECR', 'CRM-1/CRM-1 SPM-ECR'),
    
    # Coke Oven-3,4
    ('Coke Oven-3,4/CO#4 Operation Store', 'Coke Oven-3,4/CO#4 Operation Store Room'),
    ('Coke Oven-3,4/CO 3 HYDRAULIC STORE', 'Coke Oven-3,4/CO#3 Hydraulic Store room'),
    
    # WRM-1
    ('WRM-1/RGS', 'WRM-1/WRM-1 RGS'),
    
    # OBP-2
    ('OBP-2/Mechnical substores', 'OBP-2/OBP-2 Mechnical substores'),
    
    # DRI
    ('DRI/STORE', 'DRI/DRI STORE'),
    
    # Corex-1,2
    ('Corex-1,2/Mechanical substore', 'Corex-1,2/Corex -1&2 Mechanical substore'),
    ('Corex-1,2/Corex Lab', 'Technology Excellence- Iron Zone & QMC/Technology Excellence Corex Lab'),
    
    # Sinter Plant-1
    ('Sinter Plant-1/Electrical Sub Store', 'Sinter Plant-1/SP-1 Electrical Sub Store'),
    
    # SMS-2
    ('SMS-2/Ferro alloy store area', 'SMS-2/SMS-2 Ferro alloy store area'),
    
    # Research & Development
    ('Research & Development/Charazation Lab', 'Research & Development/R&D Charazation Lab'),
    
    # RMHS & Pipe Conveyor
    ('RMHS & Pipe Conveyor/ESS#4', 'RMHS & Pipe Conveyor/RMHS ESS#4'),
    ('RMHS & Pipe Conveyor/Office', 'RMHS & Pipe Conveyor/RMHS office (18 MT - JVML)'),
    
    # CRM-2
    ('CRM-2/Critical Spare Maintanance Area', 'CRM-2/CRM-2 Critical Spare Maintanance Area'),
    ('CRM-2/YMS APL (Electrical)', 'CRM-2/CRM-2 YMS - APL Electrical'),
    ('CRM-2/CRM-2 YMS ', 'CRM-2/CRM-2 YMS - APL Electrical'),  # Clean up potential half-migrated Windows folders
    ('CRM-2/CAL-Maintanance Area Spares', 'CRM-2/CRM-2 Maintanance Area Spares CAL & CGL'),
    ('CRM-2/RCL-Maintanance Area Spares', 'CRM-2/CRM-2 Maintanance Area Spares RCL'),
    ('CRM-2/Maintanance Area Spares', 'CRM-2/CRM-2 Maintanance Area Spares CAL & CGL'),
    ('CRM-2/CRM-2 Maintanance Area Spares (CRM2 RCL)', 'CRM-2/CRM-2 Maintanance Area Spares RCL'),
    ('CRM-2/CRM-2 Maintanance Area Spares (CRM2 CAL & CGL)', 'CRM-2/CRM-2 Maintanance Area Spares CAL & CGL'),
    
    # LCP-1,2,3
    ('LCP-1,2,3/Mechanical Store', 'LCP-1,2,3/LCP#2 Mechanical Store'),
    
    # BRM-2
    ('BRM-2/NTM Workshop', 'BRM-2/BRM-2 NTM Workshop'),
    
    # WRM-2
    ('WRM-2/RollShop', 'WRM-2/WRM-2 RollShop'),
    
    # Central Stores
    ('Central Stores/BAY1-CRY & CRZ Bins', 'Central Stores/Central Store BAY1-CRY & CRZ Bins')
]

for src_rel, dest_rel in folder_mappings:
    src_dir = base_assets / src_rel
    dest_dir = base_assets / dest_rel
    if src_dir.exists():
        print(f"Migrating contents of '{src_dir}' -> '{dest_dir}'")
        move_folder_contents(src_dir, dest_dir)

print("\n--- Migration completed successfully! ---")
