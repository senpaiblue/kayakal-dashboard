
import os
import io
from pathlib import Path
import re
import calendar
import datetime as dt
import base64
import csv
import uuid
from app import app
import pandas as pd
from dash import Dash, dcc, html, Input, Output, State, MATCH, ALL, ctx
import dash_bootstrap_components as dbc
import dash
from docx import Document


# ---------- CONFIG ----------
ASSETS_FOLDER = Path("./assets/5s")
EXCEL_PATH = ASSETS_FOLDER / "DT.xlsx"
DEPT_COL = "department"
MODEL_COL = "model"
LEADER_COL = "leader"

ASSETS_FOLDER.mkdir(parents=True, exist_ok=True)

# Ensure excel exists with required columns
def ensure_excel():
    if not EXCEL_PATH.exists():
        df = pd.DataFrame(columns=[DEPT_COL, MODEL_COL, LEADER_COL])
        EXCEL_PATH.parent.mkdir(parents=True, exist_ok=True)
        df.to_excel(EXCEL_PATH, index=False)

ensure_excel()

D5S_DEL_CSV = "./Data/d5s_del_requests.csv"
D5S_DEL_COLUMNS = ["id", "dept", "model", "month", "year", "type", "file", "status", "submitted_at"]

def read_d5s_del_csv():
    os.makedirs("./Data", exist_ok=True)
    if not os.path.isfile(D5S_DEL_CSV) or os.path.getsize(D5S_DEL_CSV) == 0:
        return []
    with open(D5S_DEL_CSV, newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))

def write_d5s_del_csv(rows):
    os.makedirs("./Data", exist_ok=True)
    with open(D5S_DEL_CSV, "w", newline="", encoding="utf-8") as f:
        w = csv.DictWriter(f, fieldnames=D5S_DEL_COLUMNS, extrasaction="ignore")
        w.writeheader()
        for r in rows:
            for col in D5S_DEL_COLUMNS:
                r.setdefault(col, "")
            w.writerow(r)


# ---------- 5S Audit scoring CSV (Data/5S.csv) ----------
D5S_AUDIT_CSV_PATH = Path("./Data/5S.csv")


def d5s_admin_parse_excel_bytes(file_bytes: bytes) -> pd.DataFrame:
    """Parse 5S workbook like Data/5S.xlsx → flat column names matching Data/5S.csv."""
    xl = pd.ExcelFile(io.BytesIO(file_bytes))
    sheet = "Main Score Sheet" if "Main Score Sheet" in xl.sheet_names else xl.sheet_names[0]
    df_raw = pd.read_excel(io.BytesIO(file_bytes), sheet_name=sheet, header=None)
    if df_raw.shape[0] < 5 or df_raw.shape[1] < 8:
        raise ValueError(
            "This file does not match the expected 5S score layout (need header rows and data)."
        )

    wide = df_raw.shape[1]
    score_indices = list(range(4, wide))
    last_month = None
    cols = []
    valid_indices = []
    for j in score_indices:
        m = df_raw.iloc[2, j]
        s = df_raw.iloc[3, j]
        if pd.notna(m) and str(m).strip() and str(m).lower() != "nan":
            last_month = str(m).strip().replace("Febraury", "February")
        s_clean = (
            str(s).strip()
            if pd.notna(s) and str(s).strip() and str(s).lower() != "nan"
            else ""
        )
        if not last_month or not s_clean:
            continue
        key = f"{last_month}_{s_clean}".lower().replace(" ", "_")
        cols.append(key)
        valid_indices.append(j)

    if not cols:
        raise ValueError("No score columns found (expected month + 1S…5S header row).")

    out = df_raw.iloc[4:, [1, 2, 3] + valid_indices].copy()
    if out.shape[1] != 3 + len(cols):
        raise ValueError("Column layout mismatch between header and data rows.")

    out.columns = ["sl_no", "model_area", "zone"] + cols
    out = out.dropna(how="all", subset=["model_area"])
    out = out[out["model_area"].astype(str).str.strip() != ""]
    for c in out.columns:
        out[c] = out[c].astype(str).str.strip()
    return out


def d5s_admin_merge_upload_into_csv(new_df: pd.DataFrame) -> tuple[pd.DataFrame, str]:
    """
    Merge upload into Data/5S.csv.
    Rows are keyed by (model_area, zone) case-insensitive; the latest upload wins.
    Full-row duplicates (same as existing after merge) are not written twice.
    """
    new_df = new_df.copy()
    for c in new_df.columns:
        new_df[c] = new_df[c].astype(str).str.strip()

    if D5S_AUDIT_CSV_PATH.exists():
        old_df = pd.read_csv(D5S_AUDIT_CSV_PATH, dtype=str, encoding="utf-8-sig").fillna(
            ""
        )
        for c in old_df.columns:
            old_df[c] = old_df[c].astype(str).str.strip()
    else:
        old_df = pd.DataFrame(columns=new_df.columns)

    all_cols = list(dict.fromkeys(list(old_df.columns) + list(new_df.columns)))
    old_a = old_df.reindex(columns=all_cols, fill_value="")
    new_a = new_df.reindex(columns=all_cols, fill_value="")

    combined = pd.concat([old_a, new_a], ignore_index=True)
    loc_key = (
        combined["model_area"].astype(str).str.strip().str.lower()
        + "\x1f"
        + combined["zone"].astype(str).str.strip().str.lower()
    )
    combined = combined.assign(_d5s_loc=loc_key)
    before = len(combined)
    combined = combined.drop_duplicates(subset=["_d5s_loc"], keep="last")
    after_key = len(combined)
    combined = combined.drop(columns=["_d5s_loc"])

    def _sl_sort_key(v):
        try:
            return int(float(str(v).strip()))
        except Exception:
            return 10**9

    combined = combined.sort_values(
        by="sl_no", key=lambda s: s.map(_sl_sort_key)
    ).reset_index(drop=True)

    dup_locations = before - after_key
    msg = (
        f"Saved {D5S_AUDIT_CSV_PATH} ({len(combined)} rows). "
        f"This file had {len(new_df)} row(s). "
        f"{dup_locations} row(s) at the same Model Area + Zone were merged "
        f"(only the latest upload is kept; nothing is stored twice)."
    )
    return combined, msg


# Month dropdown: 12 months starting from current month (handles year change)
import datetime as dt
import calendar

def get_month_options():
    today = dt.date.today()
    opts = []

    # from -12 months to +12 months (includes current month)
    for i in range(-12, 13):
        # calculate month & year with rollover
        month_index = today.month - 1 + i
        year = today.year + (month_index // 12)
        month_num = (month_index % 12) + 1

        label = f"{calendar.month_name[month_num]} {year}"
        value = f"{month_num}-{year}"

        opts.append({"label": label, "value": value})

    return opts


# Load unique departments from excel
def load_departments():
    try:
        df = pd.read_excel(EXCEL_PATH)
    except Exception:
        return []
    if DEPT_COL in df.columns:
        return sorted(df[DEPT_COL].dropna().astype(str).unique())
    return []

# Load model areas for given department
def load_model_areas(dept):
    if not dept:
        return []
    try:
        df = pd.read_excel(EXCEL_PATH)
    except Exception:
        return []
    if DEPT_COL not in df.columns or MODEL_COL not in df.columns:
        return []
    mask = df[DEPT_COL].astype(str).str.strip().fillna("") == str(dept).strip()
    return sorted(df.loc[mask, MODEL_COL].dropna().astype(str).unique())

# Create Dash app
INPUT_STYLE = {
    "width": "100%",
    "padding": "10px",
    "borderRadius": "10px",
    "border": "1px solid #d1d5db",
    "fontSize": "14px"
}
CARD_STYLE = {
    "background": "#ffffff",
    "borderRadius": "14px",
    "padding": "20px",
    "boxShadow": "0 10px 30px rgba(0,0,0,0.06)",
    "border": "1px solid #eef2f7"
}


mgmt_children = [
    html.Div([
        html.H2("5S Image Upload", style={
            "marginBottom": "4px",
            "fontWeight": "700"
        }),
        html.P("Upload and manage 5S images department-wise",
            style={"color": "#6b7280"})
    ], style={"marginBottom": "24px"}),
    html.Div(style={"display": "flex", "gap": "24px"}, children=[
        # LEFT: upload1 panel
        html.Div(style={**CARD_STYLE, "flex": "1"}, children=[
            html.H4("upload1 Images"),
            html.Label("Select Month (required)"),
            dcc.Dropdown(id="month-dropdown", options=get_month_options(), placeholder="Select month...", clearable=False),
            html.Br(),
            html.Label("Select Department (required)"),
            dcc.Dropdown(id="department-dropdown",
                         options=[{"label": d, "value": d} for d in load_departments()],
                         placeholder="Select department...", clearable=False),
            html.Br(),
            html.Label("Select 5S Model Area (required)"),
            dcc.Dropdown(id="model-dropdown", options=[], placeholder="Select 5S Model Area...", clearable=False),
            html.Br(),
            html.Label("Select one or more images (drag & drop)"),
            dcc.Upload(
                id="image-upload1er",
                children=html.Div([
                    html.Div("📸", style={"fontSize": "32px"}),
                    html.Div("Drag & Drop Images", style={"fontWeight": "600"}),
                    html.Div("or click to browse", style={"color": "#6b7280", "fontSize": "13px"})
                ]),
                multiple=True,
                style={
                    "width": "100%",
                    "height": "160px",
                    "border": "2px dashed #3b82f6",
                    "borderRadius": "14px",
                    "display": "flex",
                    "alignItems": "center",
                    "justifyContent": "center",
                    "flexDirection": "column",
                    "background": "#f0f7ff",
                    "cursor": "pointer",
                    "transition": "0.2s"
                }
            ),

            dcc.Store(id="selected-images-store"),

            html.Div(
                id="image-preview-container",
                style={
                    "display": "grid",
                    "gridTemplateColumns": "repeat(auto-fill, 120px)",
                    "gap": "12px",
                    "marginTop": "14px"
                }
            ),


            html.Button(
                "⬆ Upload Images",
                id="upload1-button",
                n_clicks=0,
                disabled=True,
                style={
                    "marginTop": "14px",
                    "background": "#2563eb",
                    "color": "#ffffff",
                    "border": "none",
                    "padding": "10px 20px",
                    "borderRadius": "10px",
                    "fontWeight": "600",
                    "cursor": "pointer",
                    "opacity": "0.85"
                }
            ),

            html.Div(id="upload1-output", style={"marginTop":"12px"})
        ]),

        # RIGHT: Add new Department / Model Area / Team Leader
        html.Div(style={**CARD_STYLE, "flex": "0.9"}, children=[
            html.H4("➕ Add Department / Model Area / Leader",
                style={"fontWeight": "600", "marginBottom": "16px"}),
            html.Label("Department Name"),
            dcc.Input(id="new-dept", type="text", placeholder="Enter department name", style=INPUT_STYLE),

            html.Br(), html.Br(),
            html.Label("5S Model Area"),
            dcc.Input(id="new-model", type="text", placeholder="Enter model area", style=INPUT_STYLE ),
            html.Br(), html.Br(),
            html.Label("Team Leader"),
            dcc.Input(id="new-leader", type="text", placeholder="Enter team leader name", style=INPUT_STYLE   ),
            html.Br(), html.Br(),
            html.Button(
    "💾 Save Details",
    id="save-dept-model",
    n_clicks=0,
    style={
        "background": "#10b981",
        "color": "white",
        "border": "none",
        "padding": "10px 18px",
        "borderRadius": "10px",
        "fontWeight": "600",
        "cursor": "pointer"
    }
)
,
            html.Div(id="save-output", style={"marginTop":"12px"}),
            html.Hr(),
            html.Div(id="existing-preview")
        ])
    ])
]

mgmt_children.append(
    html.Hr(style={"margin": "50px 0", "borderTop": "3px solid #2563eb"})
)

mgmt_children.append(
    dbc.Card(
        [
            dbc.CardHeader(html.H4("📝 Audit Summary Entry")),

            dbc.CardBody(
                [
                    html.Label("Department"),
                    dcc.Dropdown(
                        id="summary-department-dd",
                        options=[{"label": d, "value": d} for d in load_departments()],
                        placeholder="Select Department",
                        clearable=False,
                    ),

                    html.Br(),
                    html.Label("Model Area"),
                    dcc.Dropdown(
                        id="summary-model-dd",
                        options=[],
                        placeholder="Select Model Area",
                        clearable=False,
                    ),

                    html.Br(),
                    html.Label("Month-Year"),
                    dcc.Dropdown(
                        id="summary-month-dd",
                        options=get_month_options(),
                        placeholder="Select Month",
                        clearable=False,
                    ),

                    html.Br(),
                    html.Label("Audit Summary Text"),
                    dcc.Textarea(
                        id="summary-text",
                        style={"width": "100%", "height": "180px"},
                        placeholder="Write audit summary here..."
                    ),

                    html.Br(), html.Br(),
                    dbc.Button(
                        "💾 Save Audit Summary",
                        id="summary-save-btn",
                        color="success"
                    ),

                    html.Div(id="summary-save-msg", style={"marginTop": "12px"})
                ]
            ),
        ],
        style={
            "border": "2px solid #c7d2fe",
            "borderRadius": "14px",
            "padding": "10px",
            "background": "#ffffff"
        },
    )
)
mgmt_children.append(
    html.Hr(style={"margin": "50px 0", "borderTop": "3px solid #dc2626"})
)

mgmt_children.append(
    dbc.Card(
        [
            dbc.CardHeader(html.H4("🗑️ Admin – Delete 5S Photos")),

            dbc.CardBody(
                [
                    html.Label("Department"),
                    dcc.Dropdown(
                        id="delete-dept-dd",
                        options=[{"label": d, "value": d} for d in load_departments()],
                        clearable=False,
                    ),

                    html.Br(),
                    html.Label("Model Area"),
                    dcc.Dropdown(
                        id="delete-model-dd",
                        options=[],
                        clearable=False,
                    ),

                    html.Br(),
                    html.Label("Month-Year"),
                    dcc.Dropdown(
                        id="delete-month-dd",
                        options=get_month_options(),
                        clearable=False,
                    ),

                    html.Br(),
                    html.Div(
                        id="delete-image-preview",
                        style={
                            "display": "grid",
                            "gridTemplateColumns": "repeat(auto-fill, 140px)",
                            "gap": "14px",
                            "marginTop": "16px"
                        }
                    ),

                    html.Br(),
                    dbc.Button(
                        "🗑️ Delete Selected Photos",
                        id="delete-images-btn",
                        color="danger",
                        disabled=True
                    ),

                    html.Div(id="delete-status-msg", style={"marginTop": "12px"})
                ]
            )
        ],
        style={
            "border": "2px solid #fecaca",
            "borderRadius": "14px",
            "background": "#ffffff"
        }
    )
)

# -------- PASSWORD MANAGEMENT SECTION --------
mgmt_children.append(
    html.Hr(style={"margin": "50px 0", "borderTop": "3px solid #f59e0b"})
)

mgmt_children.append(
    dbc.Card(
        [
            dbc.CardHeader(html.H4("🔑 Department Password Management")),
            dbc.CardBody(
                [
                    html.P(
                        "Set or update the upload password for each department/model combination. "
                        "This password is required by users to unlock the 'Upload Authorization' on the 5S Self Audit page.",
                        style={"color": "#6b7280", "fontSize": "13px"}
                    ),

                    html.Label("Select Department"),
                    dcc.Dropdown(
                        id="pwd-mgmt-dept-dd",
                        options=[{"label": d, "value": d} for d in load_departments()],
                        placeholder="Select department...",
                        clearable=False,
                        style={"marginBottom": "16px"}
                    ),

                    html.Div(id="pwd-mgmt-rows-container"),

                    html.Div(id="pwd-mgmt-save-msg", style={"marginTop": "12px"})
                ]
            )
        ],
        style={
            "border": "2px solid #fde68a",
            "borderRadius": "14px",
            "background": "#fffbeb",
            "marginTop": "24px"
        }
    )
)

mgmt_children.append(
    html.Hr(style={"margin": "50px 0", "borderTop": "3px solid #22c55e"})
)

mgmt_children.append(
    dbc.Card(
        [
            dbc.CardHeader(html.H4("📊 Upload 5S Audit Scoring (Excel → CSV)")),
            dbc.CardBody(
                [
                    html.P(
                        "Upload a workbook in the same layout as Data/5S.xlsx (sheet "
                        "\"Main Score Sheet\"). It will be converted and merged into "
                        "Data/5S.csv. Rows for the same Model Area + Zone are replaced "
                        "by the new file (duplicates are not kept twice).",
                        style={"color": "#6b7280", "fontSize": "13px"},
                    ),
                    dcc.Upload(
                        id="d5s-admin-xlsx-upload",
                        children=html.Div(
                            [
                                html.Div("📄", style={"fontSize": "28px"}),
                                html.Div(
                                    "Drag & drop .xlsx here or click to browse",
                                    style={"fontWeight": "600"},
                                ),
                            ]
                        ),
                        accept=".xlsx",
                        multiple=False,
                        style={
                            "width": "100%",
                            "minHeight": "120px",
                            "border": "2px dashed #22c55e",
                            "borderRadius": "12px",
                            "display": "flex",
                            "alignItems": "center",
                            "justifyContent": "center",
                            "background": "#f0fdf4",
                            "cursor": "pointer",
                        },
                    ),
                    html.Div(id="d5s-admin-xlsx-msg", style={"marginTop": "14px"}),
                ]
            ),
        ],
        style={
            "border": "2px solid #bbf7d0",
            "borderRadius": "14px",
            "background": "#ffffff",
            "marginTop": "24px",
        },
    )
)


# Create the Tab components:
tab_5s_mgmt = dbc.Tab(
    label="5S Management",
    tab_id="tab-5s-mgmt",
    children=html.Div(mgmt_children, style={"padding": "20px"})
)

tab_del_requests = dbc.Tab(
    label="Deletion Requests",
    tab_id="tab-del-requests",
    children=dbc.Container([
        html.H4("5S Deletion Requests", className="text-center my-3 text-danger"),
        html.P("Review and approve/reject deletion requests for Audit Summary and Before/After images.", className="text-muted text-center mb-3"),
        dcc.Interval(id="d5s-del-refresh-interval", interval=5000, n_intervals=0),
        html.Div(id="d5s-del-list-container")
    ], fluid=True)
)

layout = html.Div([
    dbc.Tabs(
        [tab_5s_mgmt, tab_del_requests],
        id="d5s-admin-tabs",
        active_tab="tab-5s-mgmt",
        className="mt-3",
    )
], style={
    "fontFamily": "Inter, Segoe UI, Arial",
    "background": "linear-gradient(180deg, #eef5ff 0%, #f9fbff 100%)",
    "minHeight": "100vh",
    "padding": "24px"
})


# Callback: load rows for the selected department
@app.callback(
    Output("pwd-mgmt-rows-container", "children"),
    Input("pwd-mgmt-dept-dd", "value"),
)
def pwd_mgmt_load_rows(dept):
    if not dept:
        return html.Div("Select a department to manage passwords.", style={"color": "#9ca3af"})

    try:
        df = pd.read_excel(EXCEL_PATH, dtype=str).fillna("")
    except Exception:
        return html.Div("Could not read DT.xlsx.", style={"color": "red"})

    # Ensure password column exists
    if "password" not in df.columns:
        df["password"] = ""

    rows_df = df[df[DEPT_COL].astype(str).str.strip() == str(dept).strip()]

    if rows_df.empty:
        return html.Div("No model areas found for this department.", style={"color": "#9ca3af"})

    row_items = []
    for _, row in rows_df.iterrows():
        model_val = str(row.get(MODEL_COL, "")).strip()
        current_pwd = str(row.get("password", "")).strip()

        row_items.append(
            dbc.Row(
                [
                    dbc.Col(
                        html.Div(
                            [
                                html.Div(html.Strong(model_val), style={"fontSize": "14px"}),
                                html.Div(
                                    f"Current password: {current_pwd if current_pwd else '(none set)'}",
                                    style={"fontSize": "12px", "color": "#6b7280"}
                                )
                            ]
                        ),
                        md=4
                    ),
                    dbc.Col(
                        dbc.Input(
                            id={"type": "pwd-mgmt-input", "model": model_val},
                            type="text",
                            placeholder="Enter new password",
                            value=current_pwd,
                            style={"fontSize": "14px"}
                        ),
                        md=5
                    ),
                    dbc.Col(
                        dbc.Button(
                            "💾 Save",
                            id={"type": "pwd-mgmt-save-btn", "model": model_val},
                            color="warning",
                            size="sm",
                            n_clicks=0
                        ),
                        md=3
                    )
                ],
                className="mb-3 align-items-center"
            )
        )

    return html.Div(row_items)


# Callback: save password for a specific dept+model row
@app.callback(
    Output("pwd-mgmt-save-msg", "children"),
    Input({"type": "pwd-mgmt-save-btn", "model": dash.ALL}, "n_clicks"),
    State("pwd-mgmt-dept-dd", "value"),
    State({"type": "pwd-mgmt-input", "model": dash.ALL}, "value"),
    State({"type": "pwd-mgmt-save-btn", "model": dash.ALL}, "id"),
    prevent_initial_call=True
)
def pwd_mgmt_save_password(all_n_clicks, dept, all_pwd_values, all_btn_ids):
    ctx_trigger = dash.callback_context
    if not ctx_trigger.triggered:
        raise dash.exceptions.PreventUpdate

    # Find which button was clicked
    triggered_prop = ctx_trigger.triggered[0]["prop_id"]
    if not triggered_prop or not all_n_clicks:
        raise dash.exceptions.PreventUpdate

    # Find index of the triggered button
    triggered_idx = None
    for i, nc in enumerate(all_n_clicks):
        if nc and nc > 0:
            # check which one actually fired
            btn_str = str(all_btn_ids[i])
            if btn_str in triggered_prop:
                triggered_idx = i
                break

    if triggered_idx is None:
        # fallback: use the one with the highest n_clicks increment
        import json
        try:
            triggered_id_str = triggered_prop.split(".")[0]
            triggered_id = json.loads(triggered_id_str)
            model_key = triggered_id.get("model", "")
            triggered_idx = next(
                (i for i, bid in enumerate(all_btn_ids) if bid.get("model") == model_key),
                None
            )
        except Exception:
            raise dash.exceptions.PreventUpdate

    if triggered_idx is None:
        raise dash.exceptions.PreventUpdate

    model_to_update = all_btn_ids[triggered_idx].get("model", "")
    new_pwd = str(all_pwd_values[triggered_idx]).strip() if all_pwd_values[triggered_idx] else ""

    if not dept or not model_to_update:
        return dbc.Alert("❌ Department or Model not found.", color="danger")

    try:
        df = pd.read_excel(EXCEL_PATH, dtype=str).fillna("")
    except Exception as ex:
        return dbc.Alert(f"❌ Could not read DT.xlsx: {ex}", color="danger")

    if "password" not in df.columns:
        df["password"] = ""

    mask = (
        (df[DEPT_COL].astype(str).str.strip() == str(dept).strip()) &
        (df[MODEL_COL].astype(str).str.strip() == model_to_update)
    )

    if not mask.any():
        return dbc.Alert("❌ Row not found in Excel.", color="danger")

    df.loc[mask, "password"] = new_pwd

    try:
        df.to_excel(EXCEL_PATH, index=False)
    except Exception as ex:
        return dbc.Alert(f"❌ Could not save to Excel: {ex}", color="danger")

    return dbc.Alert(
        f"✅ Password updated for {dept} / {model_to_update}",
        color="success",
        dismissable=True
    )


@app.callback(
    Output("d5s-admin-xlsx-msg", "children"),
    Input("d5s-admin-xlsx-upload", "contents"),
    State("d5s-admin-xlsx-upload", "filename"),
    prevent_initial_call=True,
)
def d5s_admin_on_5s_xlsx_upload(contents, filename):
    if not contents:
        raise dash.exceptions.PreventUpdate

    fn = (filename or "").lower()
    if not fn.endswith(".xlsx"):
        return dbc.Alert("Please upload an .xlsx file.", color="warning", dismissable=True)

    try:
        header, b64 = contents.split(",", 1)
        raw = base64.b64decode(b64)
        new_df = d5s_admin_parse_excel_bytes(raw)
    except Exception as ex:
        return dbc.Alert(
            f"❌ Could not read Excel: {ex}",
            color="danger",
            dismissable=True,
        )

    try:
        D5S_AUDIT_CSV_PATH.parent.mkdir(parents=True, exist_ok=True)
        combined, msg = d5s_admin_merge_upload_into_csv(new_df)
        combined.to_csv(D5S_AUDIT_CSV_PATH, index=False, encoding="utf-8-sig")
    except Exception as ex:
        return dbc.Alert(
            f"❌ Could not save CSV: {ex}",
            color="danger",
            dismissable=True,
        )

    return dbc.Alert(f"✅ {msg}", color="success", dismissable=True)


# Enable upload1 button only when all required fields filled and files present
@app.callback(
    Output("upload1-button", "disabled"),
    Input("month-dropdown", "value"),
    Input("department-dropdown", "value"),
    Input("model-dropdown", "value"),
    Input("selected-images-store", "data"),
)
def toggle_upload1_disabled(month_val, dept, model, images):
    return not (month_val and dept and model and images)


# Populate model dropdown based on department
@app.callback(
    Output("model-dropdown", "options"),
    Input("department-dropdown", "value")
)
def update_model_options(dept):
    areas = load_model_areas(dept)
    return [{"label": a, "value": a} for a in areas]

# Handle upload1s
@app.callback(
    Output("upload1-output", "children"),
    Input("upload1-button", "n_clicks"),   # ✅ REQUIRED
    State("selected-images-store", "data"),
    State("month-dropdown", "value"),
    State("department-dropdown", "value"),
    State("model-dropdown", "value"),
    prevent_initial_call=True
)
def handle_upload1(n_clicks, stored_images, month_val, dept, model):

    if not stored_images:
        return html.Div(
            "Error: All fields mandatory and at least one image must be selected.",
            style={"color": "red"}
        )

    # ---------- PARSE MONTH ----------
    try:
        m_str, y_str = month_val.split("-")
        month_num = int(m_str)
        year_num = int(y_str)
        month_name = calendar.month_name[month_num].lower()
    except Exception as e:
        return html.Div(f"Invalid month selection: {e}", style={"color": "red"})

    # ---------- TARGET FOLDER ----------
    folder = ASSETS_FOLDER / str(dept).strip() / str(model).strip()
    folder.mkdir(parents=True, exist_ok=True)

    # ---------- FIND START INDEX ----------
    existing_indices = []
    for f in folder.iterdir():
        if f.is_file():
            stem = f.stem.lower()
            m = re.search(
                rf"^(\d+).*{re.escape(month_name)}{year_num}",
                stem
            )
            if m:
                try:
                    existing_indices.append(int(m.group(1)))
                except Exception:
                    pass

    current_index = max(existing_indices) + 1 if existing_indices else 1

    # ---------- SAVE FILES ----------
    saved = []

    for item in stored_images:
        try:
            content = item["content"]
            filename = item.get("filename")

            header, b64 = content.split(",", 1)
            ext = Path(filename).suffix if filename else ".jpg"
            if not ext:
                ext = ".jpg"

            save_name = f"{current_index}{month_name}{year_num}{ext}"
            save_path = folder / save_name

            with open(save_path, "wb") as fh:
                fh.write(base64.b64decode(b64))

            saved.append(save_name)
            current_index += 1

        except Exception:
            continue  # skip broken image but continue

    # ---------- RESULT ----------
    if not saved:
        return html.Div(
            "No images saved (something went wrong).",
            style={"color": "orange"}
        )

    return dbc.Alert(
        [
            html.Div(f"✅ Saved {len(saved)} image(s) successfully"),
            html.Ul([html.Li(s) for s in saved])
        ],
        color="success",
        dismissable=True
    )


# Save new Department + ModelArea + Team Leader (prevents duplicate dept+model)
@app.callback(
    Output("save-output", "children"),
    Output("department-dropdown", "options"),
    Output("existing-preview", "children"),
    Input("save-dept-model", "n_clicks"),
    State("new-dept", "value"),
    State("new-model", "value"),
    State("new-leader", "value")
)
def save_dept_model(n_clicks, new_dept, new_model, new_leader):
    # initial load: show existing departments
    if not n_clicks:
        depts = load_departments()
        preview = html.Div([html.H5("Existing Departments"), html.Ul([html.Li(d) for d in depts])])
        return dash.no_update, [{"label": d, "value": d} for d in depts], preview

    # validations
    if not new_dept or not new_model or not new_leader:
        return html.Div("Please fill all fields.", style={"color":"red"}), dash.no_update, dash.no_update

    new_dept = str(new_dept).strip()
    new_model = str(new_model).strip()
    new_leader = str(new_leader).strip()

    # load df and check duplicates
    try:
        df = pd.read_excel(EXCEL_PATH)
    except Exception:
        df = pd.DataFrame(columns=[DEPT_COL, MODEL_COL, LEADER_COL])

    # ensure required columns exist
    for c in [DEPT_COL, MODEL_COL, LEADER_COL]:
        if c not in df.columns:
            df[c] = ""

    dup_mask = (df[DEPT_COL].astype(str).str.strip() == new_dept) & (df[MODEL_COL].astype(str).str.strip() == new_model)
    if dup_mask.any():
        return html.Div("This Department + Model Area already exists!", style={"color":"orange"}), dash.no_update, dash.no_update

    # append row and save
    df = pd.concat([df, pd.DataFrame([{DEPT_COL: new_dept, MODEL_COL: new_model, LEADER_COL: new_leader}])], ignore_index=True)
    try:
        df.to_excel(EXCEL_PATH, index=False)
    except Exception as e:
        return html.Div(f"Failed to save to Excel: {e}", style={"color":"red"}), dash.no_update, dash.no_update

    depts = load_departments()
    preview = html.Div([html.H5("Existing Departments"), html.Ul([html.Li(d) for d in depts])])
    return html.Div("Saved successfully!", style={"color":"green"}), [{"label": d, "value": d} for d in depts], preview



@app.callback(
    Output("image-preview-container", "children"),
    Output("selected-images-store", "data"),
    Input("image-upload1er", "contents"),
    Input({"type": "remove-image", "index": dash.ALL}, "n_clicks"),
    State("selected-images-store", "data"),
    State("image-upload1er", "filename"),
    prevent_initial_call=True
)
def manage_image_preview(contents, remove_clicks, stored, filenames):

    ctx = dash.callback_context
    trigger = ctx.triggered_id

    # ---------- FIRST UPLOAD ----------
    if trigger == "image-upload1er":
        if not contents:
            return [], []

        stored = []
        for i, content in enumerate(contents):
            stored.append({
                "content": content,
                "filename": filenames[i]
            })

    # ---------- REMOVE IMAGE ----------
    elif isinstance(trigger, dict) and trigger.get("type") == "remove-image":
        idx = trigger["index"]
        if stored and idx < len(stored):
            stored.pop(idx)

    # ---------- BUILD PREVIEWS ----------
    previews = []
    for i, item in enumerate(stored):
        previews.append(
            html.Div(
                style={
                    "position": "relative",
                    "border": "1px solid #ddd",
                    "borderRadius": "10px",
                    "padding": "6px",
                    "background": "#fff"
                },
                children=[
                    html.Img(
                        src=item["content"],
                        style={"width": "100%", "borderRadius": "6px"}
                    ),
                    html.Button(
                        "✖",
                        id={"type": "remove-image", "index": i},
                        style={
                            "position": "absolute",
                            "top": "4px",
                            "right": "4px",
                            "border": "none",
                            "background": "#dc3545",
                            "color": "white",
                            "borderRadius": "50%",
                            "width": "22px",
                            "height": "22px",
                            "cursor": "pointer"
                        }
                    )
                ]
            )
        )

    return previews, stored

@app.callback(
    Output("summary-model-dd", "options"),
    Output("summary-model-dd", "value"),
    Input("summary-department-dd", "value"),
)
def update_summary_model_dropdown(dept):
    if not dept:
        return [], None

    models = load_model_areas(dept)
    return [{"label": m, "value": m} for m in models], None
@app.callback(
    Output("summary-save-msg", "children"),
    Input("summary-save-btn", "n_clicks"),
    State("summary-department-dd", "value"),
    State("summary-model-dd", "value"),
    State("summary-month-dd", "value"),
    State("summary-text", "value"),
    prevent_initial_call=True
)
def save_audit_summary(n, dept, model, month_val, text):

    if not dept or not model:
        return dbc.Alert("❌ Select Department and Model Area.", color="danger")

    if not month_val:
        return dbc.Alert("❌ Select Month.", color="danger")

    if not text or not text.strip():
        return dbc.Alert("❌ Summary text cannot be empty.", color="danger")

    # Parse month
    try:
        m_str, y_str = month_val.split("-")
        month_num = int(m_str)
        year = int(y_str)
        month_name = calendar.month_name[month_num].lower()
    except Exception:
        return dbc.Alert("❌ Invalid month format.", color="danger")

    # Save path (same folder structure as images)
    folder = ASSETS_FOLDER / str(dept).strip() / str(model).strip()
    folder.mkdir(parents=True, exist_ok=True)

    filename = f"{month_name}{year}.docx"   # december2025.docx
    file_path = folder / filename

    doc = Document()
    #doc.add_heading(f"Audit Summary – {month_name.capitalize()} {year}", level=1)

    for line in text.splitlines():
        doc.add_paragraph(line)

    doc.save(file_path)

    return dbc.Alert(
        f"✅ Audit summary saved as {filename}",
        color="success",
        dismissable=True
    )
@app.callback(
    Output("delete-model-dd", "options"),
    Output("delete-model-dd", "value"),
    Input("delete-dept-dd", "value"),
)
def update_delete_model_dropdown(dept):
    if not dept:
        return [], None

    models = load_model_areas(dept)
    return [{"label": m, "value": m} for m in models], None
@app.callback(
    Output("delete-image-preview", "children"),
    Input("delete-dept-dd", "value"),
    Input("delete-model-dd", "value"),
    Input("delete-month-dd", "value"),
)
def load_images_for_deletion(dept, model, month_val):

    if not dept or not model or not month_val:
        return []

    try:
        m_str, y_str = month_val.split("-")
        month_num = int(m_str)
        year = int(y_str)
        month_name = calendar.month_name[month_num].lower()
    except Exception:
        return []

    folder = ASSETS_FOLDER / str(dept).strip() / str(model).strip()
    if not folder.exists():
        return []

    cards = []

    for img in sorted(folder.iterdir()):
        if img.is_file() and month_name in img.stem.lower() and str(year) in img.stem:
            encoded = base64.b64encode(img.read_bytes()).decode()

            cards.append(
                html.Div(
                    [
                        dcc.Checklist(
                            options=[{"label": "", "value": str(img)}],
                            id={"type": "delete-checkbox", "index": str(img)},
                            style={"position": "absolute", "top": "6px", "left": "6px"}
                        ),
                        html.Img(
                            src=f"data:image/jpeg;base64,{encoded}",
                            style={
                                "width": "100%",
                                "borderRadius": "10px",
                                "border": "1px solid #ddd"
                            }
                        )
                    ],
                    style={
                        "position": "relative",
                        "padding": "6px",
                        "background": "#f9fafb",
                        "borderRadius": "12px"
                    }
                )
            )

    return cards
@app.callback(
    Output("delete-images-btn", "disabled"),
    Input({"type": "delete-checkbox", "index": dash.ALL}, "value"),
)
def toggle_delete_button(values):
    return not any(values)
@app.callback(
    Output("delete-status-msg", "children"),
    Input("delete-images-btn", "n_clicks"),
    State({"type": "delete-checkbox", "index": dash.ALL}, "value"),
    prevent_initial_call=True
)
def delete_selected_images(n, values):

    deleted = []

    for v in values:
        if v:
            try:
                path = Path(v[0])
                if path.exists():
                    path.unlink()
                    deleted.append(path.name)
            except Exception:
                continue

    if not deleted:
        return dbc.Alert("No images deleted.", color="warning")

    return dbc.Alert(
        [
            html.Div(f"✅ Deleted {len(deleted)} image(s)"),
            html.Ul([html.Li(name) for name in deleted])
        ],
        color="success",
        dismissable=True
    )


# ----------------- 5S DELETION REQUESTS HANDLERS -----------------

IMG_FRAME_D5S = {
    "width": "100%",
    "height": "200px",
    "backgroundColor": "#111",
    "border": "1px solid #ccc",
    "display": "flex",
    "alignItems": "center",
    "justifyContent": "center",
    "overflow": "hidden",
}
IMG_STYLE_D5S = {
    "maxWidth": "100%",
    "maxHeight": "100%",
    "objectFit": "contain",
}

def encode_image_to_datauri(path):
    if not os.path.exists(path):
        return ""
    with open(path, "rb") as f:
        data = base64.b64encode(f.read()).decode()
    return f"data:image/jpg;base64,{data}"

def _build_d5s_del_card(row):
    uid = row["id"]
    dept = row.get("dept", "")
    model = row.get("model", "")
    month = row.get("month", "")
    year = row.get("year", "")
    req_type = row.get("type", "")
    filename = row.get("file", "")
    submitted_at = row.get("submitted_at", "")
    
    folder = os.path.join("assets", "5s", str(dept), str(model))
    file_path = os.path.join(folder, filename)
    
    preview_content = None
    if req_type == "summary":
        if os.path.exists(file_path):
            try:
                doc = Document(file_path)
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                preview_content = html.Div(
                    text,
                    style={
                        "whiteSpace": "pre-line",
                        "maxHeight": "200px",
                        "overflowY": "auto",
                        "border": "1px solid #ccc",
                        "padding": "10px",
                        "backgroundColor": "#f8f9fa",
                        "borderRadius": "4px",
                        "fontSize": "13px"
                    }
                )
            except Exception as e:
                preview_content = html.Div(f"Error reading summary docx: {e}", style={"color": "red"})
        else:
            preview_content = html.Div("(Summary docx file not found / already deleted)", style={"color": "orange", "fontStyle": "italic"})
    else:
        if os.path.exists(file_path):
            img_uri = encode_image_to_datauri(file_path)
            preview_content = html.Div(
                html.Img(src=img_uri, style=IMG_STYLE_D5S),
                style=IMG_FRAME_D5S,
            )
        else:
            preview_content = html.Div("(Image file not found / already deleted)", style={"color": "orange", "fontStyle": "italic"})
            
    badge_color = "primary" if req_type == "before" else ("info" if req_type == "after" else "warning")
    
    return dbc.Card(
        dbc.CardBody([
            dbc.Row([
                dbc.Col([
                    html.Small(f"5S Deletion — {req_type.capitalize()}", className="fw-bold text-muted d-block mb-1"),
                    preview_content
                ], md=6),
                dbc.Col([
                    html.Div([
                        html.Div([
                            html.Span("Type: ", className="fw-bold"),
                            dbc.Badge(req_type.capitalize(), color=badge_color, className="ms-1"),
                        ]),
                        html.Div([
                            html.Span("Department: ", className="fw-bold"),
                            html.Span(dept),
                        ], className="mt-1"),
                        html.Div([
                            html.Span("Model Area: ", className="fw-bold"),
                            html.Span(model),
                        ], className="mt-1"),
                        html.Div([
                            html.Span("Month/Year: ", className="fw-bold"),
                            html.Span(f"{month.capitalize()} {year}"),
                        ], className="mt-1"),
                        html.Div([
                            html.Span("File: ", className="fw-bold"),
                            html.Span(filename, className="text-break"),
                        ], className="mt-1"),
                        html.Div([
                            html.Span("Submitted: ", className="fw-bold"),
                            html.Span(submitted_at),
                        ], className="text-muted small mt-1"),
                        
                        html.Div([
                            dbc.Button(
                                "✔ Approve Deletion",
                                id={"type": "d5s-del-approve-btn", "uid": uid},
                                color="success",
                                size="sm",
                                className="me-2",
                            ),
                            dbc.Button(
                                "✘ Reject",
                                id={"type": "d5s-del-reject-btn", "uid": uid},
                                color="danger",
                                size="sm",
                            ),
                        ], className="mt-3"),
                        html.Div(id={"type": "d5s-del-result", "uid": uid})
                    ])
                ], md=6)
            ])
        ]),
        className="mb-3 shadow-sm",
        style={"border": "1px solid #eef2f7", "borderRadius": "10px"}
    )


@app.callback(
    Output("d5s-del-list-container", "children"),
    Input("d5s-del-refresh-interval", "n_intervals"),
    prevent_initial_call=False,
)
def d5s_del_refresh_list(_interval_tick):
    """Load all pending deletion requests from d5s_del_requests.csv."""
    cards = []
    for r in read_d5s_del_csv():
        if r.get("status") == "pending":
            cards.append(_build_d5s_del_card(r))

    if not cards:
        return dbc.Alert("No pending deletion requests 🎉", color="info")
    return cards


@app.callback(
    Output({"type": "d5s-del-result", "uid": MATCH}, "children"),
    Input({"type": "d5s-del-approve-btn", "uid": MATCH}, "n_clicks"),
    prevent_initial_call=True,
)
def d5s_del_approve(_):
    uid = ctx.triggered_id["uid"]
    all_rows = read_d5s_del_csv()
    
    target = None
    for r in all_rows:
        if r["id"] == uid:
            target = r
            break
            
    if not target:
        return dbc.Alert("Request not found", color="warning")
        
    dept = target.get("dept", "")
    model = target.get("model", "")
    filename = target.get("file", "")
    
    folder = os.path.join("assets", "5s", str(dept), str(model))
    file_path = os.path.join(folder, filename)
    
    if os.path.isfile(file_path):
        try:
            os.remove(file_path)
        except Exception as e:
            return dbc.Alert(f"Error deleting file: {e}", color="danger", className="mt-2 py-1")
            
    # Mark as approved
    for r in all_rows:
        if r["id"] == uid:
            r["status"] = "approved"
            
    write_d5s_del_csv(all_rows)
    return dbc.Alert("Deleted completely ✔", color="success", className="mt-2 py-1")


@app.callback(
    Output({"type": "d5s-del-result", "uid": MATCH}, "children", allow_duplicate=True),
    Input({"type": "d5s-del-reject-btn", "uid": MATCH}, "n_clicks"),
    prevent_initial_call=True,
)
def d5s_del_reject(_):
    uid = ctx.triggered_id["uid"]
    all_rows = read_d5s_del_csv()
    
    target = None
    for r in all_rows:
        if r["id"] == uid:
            target = r
            break
            
    if not target:
        return dbc.Alert("Request not found", color="warning")
        
    # Mark as rejected — leave file untouched
    for r in all_rows:
        if r["id"] == uid:
            r["status"] = "rejected"
            
    write_d5s_del_csv(all_rows)
    return dbc.Alert("Rejected — file kept ✘", color="danger", className="mt-2 py-1")

